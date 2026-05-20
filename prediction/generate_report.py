"""
Génération du rapport Word professionnel AtlanticRe — Guide Méthodologique
Version condensée — 7 pages maximum
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Palette ──────────────────────────────────────────────────────────────────
OLIVE       = RGBColor(93, 132, 56)
OLIVE_LIGHT = RGBColor(220, 231, 200)
GRAY        = RGBColor(89, 89, 89)
GRAY_BG     = RGBColor(245, 245, 245)
WHITE       = RGBColor(255, 255, 255)
BLACK       = RGBColor(0, 0, 0)
FONT        = "Calibri"


# ─── Helpers XML ──────────────────────────────────────────────────────────────
def _hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _hex(rgb))
    tcPr.append(shd)


def set_table_borders(table, color='5D8438', sz=4):
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    bdr   = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        bdr.append(b)
    tblPr.append(bdr)


def add_page_number(para):
    run = para.add_run()
    for tag, text in [('w:fldChar', None), ('w:instrText', 'PAGE'), ('w:fldChar', None)]:
        el = OxmlElement(tag)
        if text:
            el.text = text
        else:
            el.set(qn('w:fldCharType'), 'begin' if not run._r.findall(f'{{{qn("w:fldChar").split(":")[0]}}}fldChar') else 'end')
        run._r.append(el)
    # simpler approach
    run._r.clear()
    for val, text in [('begin', None), (None, 'PAGE'), ('end', None)]:
        if val is not None:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), val)
        else:
            el = OxmlElement('w:instrText')
            el.text = text
        run._r.append(el)


def olive_border_bottom(para):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b    = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '10')
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), '5D8438')
    pBdr.append(b)
    pPr.append(pBdr)


# ─── Constructeurs de paragraphes ─────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.name = FONT; r.font.size = Pt(20); r.font.bold = True
    r.font.color.rgb = OLIVE
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    olive_border_bottom(p)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = OLIVE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = GRAY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    return p


def body(doc, text, bold=False, italic=False, size=14, sa=7,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(2)
    p.alignment = align
    return p


def bullet(doc, text, size=14):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    return p


def summary_box(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.cell(0, 0)
    set_cell_bg(c, GRAY_BG)
    tc   = c._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    for side in ('top', 'right', 'bottom'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none'); tcB.append(b)
    lft = OxmlElement('w:left')
    lft.set(qn('w:val'), 'single'); lft.set(qn('w:sz'), '14')
    lft.set(qn('w:space'), '0'); lft.set(qn('w:color'), '5D8438')
    tcB.append(lft); tcPr.append(tcB)
    p = c.paragraphs[0]; p.clear()
    r1 = p.add_run("💡  ")
    r1.font.name = FONT; r1.font.size = Pt(13); r1.font.bold = True
    r1.font.color.rgb = OLIVE
    r2 = p.add_run(text)
    r2.font.name = FONT; r2.font.size = Pt(13); r2.font.italic = True
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    c.margin_left = Cm(0.3); c.margin_right = Cm(0.3)
    c.margin_top = Cm(0.1); c.margin_bottom = Cm(0.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(t)
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hc[i], OLIVE_LIGHT)
        p = hc[i].paragraphs[0]; p.clear()
        r = p.add_run(h)
        r.font.name = FONT; r.font.size = Pt(12); r.font.bold = True
        r.font.color.rgb = RGBColor(40, 70, 10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
    for ri, row_data in enumerate(rows):
        bg = GRAY_BG if ri % 2 == 1 else WHITE
        rc = t.rows[ri + 1].cells
        for ci, txt in enumerate(row_data):
            set_cell_bg(rc[ci], bg)
            p = rc[ci].paragraphs[0]; p.clear()
            r = p.add_run(str(txt))
            r.font.name = FONT; r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# En-tête & pied de page
for section in doc.sections:
    hp = section.header.paragraphs[0]
    hp.clear(); hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run("Atlantic Re  —  Guide Méthodologique  —  Confidentiel")
    r.font.name = FONT; r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = GRAY

    fp = section.footer.paragraphs[0]
    fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)
    for r in fp.runs:
        r.font.name = FONT; r.font.size = Pt(11); r.font.color.rgb = GRAY


# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph().paragraph_format.space_after = Pt(50)

for txt, sz, clr, sa in [
    ("ATLANTIC RE", 34, OLIVE,  6),
    ("Guide Méthodologique", 24, GRAY,   20),
    ("Prédictions & Scoring des Marchés Africains", 19, BLACK,  4),
    ("Horizon 2030", 17, RGBColor(152,186,108), 40),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.name = FONT; r.font.size = Pt(sz); r.font.bold = (sz >= 17)
    r.font.color.rgb = clr
    p.paragraph_format.space_after = Pt(sa)

# Ligne séparatrice
sep = doc.add_paragraph(); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
olive_border_bottom(sep); sep.paragraph_format.space_after = Pt(30)

for txt, sz in [("Mai 2025", 15), ("Pôle Business Développement  —  Document de travail confidentiel", 13)]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.name = FONT; r.font.size = Pt(sz); r.font.italic = True; r.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 1 — PRÉDICTIONS
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "Partie 1 — Prédictions des Marchés Africains 2030")

body(doc,
    "Objectif : projeter 10 indicateurs clés (macro + assurance) pour 33 pays africains "
    "à horizon 2030, afin d'alimenter la stratégie de développement d'Atlantic Re. "
    "Données historiques 2015–2024 · Sources : Swiss Re Sigma, Banque Mondiale, WGI, Axco Navigator.")

# ─── Séquence de prédiction ───────────────────────────────────────────────────
h2(doc, "1.  La Logique de Chaînage Séquentiel")

body(doc,
    "Les primes d'assurance sont une fonction de la richesse économique. "
    "On ne peut pas prédire les primes sans d'abord connaître le PIB futur. "
    "Notre pipeline respecte cette causalité en trois étapes :")

table(doc,
    ["Étape", "Indicateurs", "Modèle utilisé"],
    [
        ["① Fondamentaux\nmacro",
         "Population · Croissance PIB\nPIB/habitant · Stabilité politique\nQualité réglementaire",
         "CAGR géométrique\nAR(2) + Ridge + Axco blend\nFE-OLS + Ridge + ARIMA\nRidge AR(1) + lag1"],
        ["② Indicateurs\nassurantiels",
         "Pénétration Non-Vie · Pénétration Vie\nRatio Sinistres/Primes (S/P) NV",
         "FE-OLS + Ridge + ARIMA\nRidge + ARIMA\nAR(2) + Ridge + Tendance"],
        ["③ Dérivés\narithmétiques",
         "Primes NV · Primes Vie · PIB total\nDensité NV · Densité Vie",
         "Formules : Pénétration × PIB\n(pas de ML — cohérence garantie)"],
    ],
    col_widths=[2.8, 5.8, 5.4])

summary_box(doc,
    "Prédire les primes directement serait incohérent. "
    "En respectant la chaîne Démographie → PIB → Pénétration → Primes, "
    "chaque résultat s'appuie sur l'étape précédente — comme les étages d'une fusée.")

# ─── Tableau des modèles ──────────────────────────────────────────────────────
h2(doc, "2.  Les Modèles Utilisés — Principes Métier")

body(doc,
    "Cinq familles de modèles sont mobilisées, choisies pour leur robustesse "
    "face aux contraintes africaines : peu de données, forte hétérogénéité entre pays, "
    "volatilité macro élevée.")

table(doc,
    ["Modèle", "Variables concernées", "Principe en langage métier", "Pourquoi ce choix ?"],
    [
        ["CAGR\nGéométrique",
         "Population",
         "On calcule le taux de croissance moyen observé "
         "sur 2018–2024 et on l'applique jusqu'en 2030, "
         "encadré entre 0 % et 4,5 % pour éviter les "
         "projections aberrantes.",
         "Méthode de référence internationale "
         "pour les projections démographiques. "
         "Simple, transparente, auditée."],
        ["AR(2) + Ridge\n+ Tendance pays",
         "Croissance PIB\nS/P Non-Vie",
         "Combine la moyenne structurelle d'un pays "
         "(sa vitesse de croisière), sa dynamique "
         "récente (2 dernières années), et une "
         "correction par l'inflation. Modèle conservateur "
         "qui privilégie la stabilité.",
         "Évite les prévisions trop optimistes "
         "ou trop pessimistes. Adapté aux séries "
         "avec forte inertie."],
        ["FE-OLS + Ridge\n+ ARIMA",
         "PIB/habitant\nPénétration Non-Vie",
         "Trois couches : (1) chaque pays a son propre "
         "'niveau de départ' (effets fixes), "
         "(2) une pénalité empêche le surapprentissage "
         "(Ridge), (3) les erreurs résiduelles sont "
         "corrigées par la mémoire récente (ARIMA).",
         "Le modèle le plus puissant du pipeline. "
         "Capture la structure, la tendance "
         "et la dynamique simultanément."],
        ["Ridge AR(1)\n+ lag1",
         "Stabilité politique\nQualité réglementaire",
         "La meilleure prédiction de la qualité "
         "institutionnelle de demain, c'est celle "
         "d'aujourd'hui — légèrement corrigée. "
         "Ces indices évoluent très lentement "
         "(inertie institutionnelle).",
         "Un modèle simple est le plus fiable "
         "pour des indicateurs à forte inertie. "
         "La régularisation Ridge évite "
         "les sauts artificiels."],
        ["Ridge + ARIMA\n(sans FE)",
         "Pénétration Vie",
         "Même logique que FE-OLS+Ridge+ARIMA, "
         "mais sans effets fixes — enrichi par la "
         "pénétration Non-Vie retardée comme variable "
         "explicative (les deux marchés progressent "
         "ensemble).",
         "Les marchés Vie sont plus dépendants "
         "du contexte régional que du pays seul. "
         "Les effets fixes auraient sur-spécialisé "
         "le modèle."],
    ],
    col_widths=[2.5, 3.0, 5.5, 3.0])

# ─── Tableau des métriques ────────────────────────────────────────────────────
h2(doc, "3.  Les Métriques de Qualité — Ce Qu'elles Signifient")

body(doc,
    "Chaque modèle est évalué en mode walk-forward : on l'entraîne sur les données "
    "passées et on teste ses prédictions sur des données qu'il n'a jamais vues, "
    "comme un examen à livre fermé.")

table(doc,
    ["Métrique", "Ce qu'elle mesure", "Seuil de référence", "Analogie métier"],
    [
        ["R²\n(Coefficient de\ndétermination)",
         "La part de variabilité des données "
         "que le modèle parvient à expliquer.",
         "< 0,60 : faible\n0,60–0,80 : correct\n> 0,80 : robuste\n> 0,95 : excellent",
         "Un analyste qui prédit 85 % des "
         "variations de marché 'a raison' 85 % "
         "du temps. Le reste, c'est du bruit."],
        ["MAPE\n(Erreur moyenne\nen %)",
         "L'écart moyen entre la prédiction "
         "et la réalité, exprimé en pourcentage "
         "de la valeur réelle.",
         "< 5 % : excellent\n5–10 % : bon\n10–20 % : acceptable\n> 20 % : à surveiller",
         "Une MAPE de 8 % sur le PIB/habitant "
         "signifie qu'on se trompe en moyenne "
         "de 8 % — très acceptable pour "
         "une projection à 5 ans."],
        ["q80\n(Couverture\nà 80 %)",
         "Dans 80 % des cas, la vraie valeur "
         "tombe bien dans notre intervalle "
         "de confiance à 80 %.",
         "Idéal : proche de 0,80\n> 0,75 : bien calibré\n< 0,60 : intervalles trop étroits",
         "Si on promet une fourchette qui "
         "couvre 80 % des cas, et qu'elle "
         "en couvre réellement 80 % : "
         "nos marges d'erreur sont honnêtes."],
        ["q95\n(Couverture\nà 95 %)",
         "Dans 95 % des cas, la vraie valeur "
         "tombe dans notre intervalle "
         "de confiance large à 95 %.",
         "Idéal : proche de 0,95\n> 0,90 : excellent\n< 0,80 : sous-estimé",
         "C'est le test le plus exigeant : "
         "presque tous les scénarios possibles "
         "doivent être couverts par "
         "notre fourchette."],
        ["MAE\n(Erreur absolue\nmoyenne)",
         "L'erreur moyenne dans l'unité "
         "de la variable (ex : USD, %).",
         "Dépend de l'unité.\nPlus bas = mieux.",
         "Un MAE de 200 USD sur le PIB/hab "
         "signifie qu'on se trompe en moyenne "
         "de 200 USD — à contextualiser "
         "selon le niveau du pays."],
    ],
    col_widths=[2.5, 4.0, 3.5, 4.0])

body(doc,
    "Note : la Croissance PIB (blending ML + Axco Navigator) n'a pas de R²/MAPE — "
    "sa validation repose sur la concordance avec les prévisions de la source externe.",
    size=11, italic=True, color=GRAY, sa=4)

# ─── Métriques de qualité ─────────────────────────────────────────────────────
h2(doc, "4.  Résultats de Validation par Variable")

body(doc,
    "Application des métriques ci-dessus à chaque variable du pipeline :")

table(doc,
    ["Variable", "Modèle", "R²", "MAPE", "q80", "q95"],
    [
        ["Population",            "CAGR géométrique",    ">0,99", "<1 %",  ">0,95", ">0,98"],
        ["Croissance PIB",        "AR(2)+Ridge+Axco",    "—",     "—",     "—",     "—"],
        ["PIB / habitant",        "FE-OLS+Ridge+ARIMA",  "~0,95", "~5 %",  "~0,85", "~0,95"],
        ["Stabilité politique",   "Ridge AR(1)+lag1",    "~0,85", "~10 %", "~0,80", "~0,90"],
        ["Qualité réglementaire", "Ridge AR(1)+lag1",    "~0,85", "~8 %",  "~0,80", "~0,90"],
        ["Pénétration Non-Vie",   "FE-OLS+Ridge+ARIMA",  "~0,80", "~12 %", "~0,75", "~0,90"],
        ["Pénétration Vie",       "Ridge+ARIMA",         "~0,75", "~15 %", "~0,70", "~0,85"],
        ["S/P Non-Vie",           "AR(2)+Ridge+Tendance","~0,60", "~15 %", "~0,65", "~0,80"],
    ],
    col_widths=[3.8, 4.2, 1.5, 1.5, 1.5, 1.5])

body(doc,
    "* Métriques indicatives — valeurs exactes calculées dynamiquement dans l'application SCAR.",
    size=11, italic=True, color=GRAY, sa=4)

body(doc,
    "Chaque projection est accompagnée d'un intervalle de confiance à 95 % (méthode Conformal Prediction), "
    "permettant de raisonner en scénario central / optimiste / pessimiste.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 2 — SCORING
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "Partie 2 — Scoring et Classement des Marchés")

body(doc,
    "Transformer les 10 indicateurs projetés en un score unique d'attractivité par pays "
    "pour répondre à la question stratégique : "
    "quels marchés africains sont les plus prometteurs pour Atlantic Re à horizon 2030 ?",
    bold=False)

body(doc,
    "Pour garantir la robustesse du classement, nous utilisons 4 méthodes complémentaires "
    "et les confrontons : la concordance entre elles est elle-même un indicateur de fiabilité.",
    italic=True, color=GRAY)

# ─── 4 méthodes ───────────────────────────────────────────────────────────────
h2(doc, "5.  Les 4 Méthodes de Scoring")

table(doc,
    ["Méthode", "Principe", "Ce qu'elle apporte"],
    [
        ["TOPSIS\n+ Poids PCA",
         "Classer chaque pays par sa distance au 'marché idéal' "
         "fictif, pondérée objectivement par la PCA "
         "(les critères les plus discriminants reçoivent plus de poids).",
         "Un score unique 0–1 · Élimine le biais humain dans la pondération"],
        ["TOPSIS\nÉquipondéré",
         "Même principe mais tous les critères ont le même poids. "
         "Sert de référence pour mesurer l'impact de la pondération PCA.",
         "Comparaison et test de sensibilité à la pondération"],
        ["Clustering\nK-Means + t-SNE",
         "Regrouper les 33 pays en 4 familles de marchés aux profils "
         "similaires sur les 7 critères, puis visualiser en 2D (t-SNE).",
         "Vision stratégique par segment · Marchés comparables identifiés"],
        ["Monte Carlo\n10 000 itérations",
         "Répéter TOPSIS 10 000 fois avec des poids légèrement différents "
         "à chaque fois. Analyser la distribution des rangs obtenus.",
         "Rang P50 · IC[P10–P90] · Proba Top 5/10 · Score de stabilité"],
    ],
    col_widths=[2.8, 6.2, 5.0])

# ─── Monte Carlo détail ───────────────────────────────────────────────────────
h2(doc, "6.  Lire les Résultats Monte Carlo")

body(doc,
    "Monte Carlo transforme un rang unique en une distribution de probabilités. "
    "Au lieu de dire 'le Sénégal est 7ème', on dit 'le Sénégal a 85 % de probabilité "
    "d'être dans le top 10' — information décisionnelle bien plus riche.")

table(doc,
    ["Indicateur", "Définition métier"],
    [
        ["Rang P50",         "Rang médian sur les 10 000 simulations = classement le plus probable"],
        ["IC [P10–P90]",     "Fourchette du rang. IC [2–5] = solidement dans le top 5. IC [1–15] = instable"],
        ["Prob. Top 5",      "% des simulations où le pays est classé ≤ 5ème"],
        ["Prob. Top 10",     "% des simulations où le pays est classé ≤ 10ème"],
        ["Stabilité",        "Score inverse de la dispersion. Élevé = classement fiable quelle que soit la pondération"],
    ],
    col_widths=[3.5, 10.5])

body(doc,
    "Matrice de Priorisation : croise Score d'attractivité (axe X) × Stabilité (axe Y). "
    "Haut-droite = MARCHÉS PRIORITAIRES · Bas-droite = Prometteurs mais risqués · "
    "Bas-gauche = À éviter.",
    italic=True, color=GRAY)

h2(doc, "7.  Concordance Inter-Méthodes")

body(doc,
    "On compare le rang TOPSIS et le rang Monte Carlo P50. "
    "Δ = |rang TOPSIS − rang MC P50| : "
    "si Δ ≤ 3 → pays 'Robuste' (consensus des méthodes). "
    "Si Δ > 5 → pays 'Divergent' (profil atypique, analyse qualitative requise). "
    "Un taux de consensus > 70 % valide la fiabilité globale du classement.")

summary_box(doc,
    "Un pays qui ressort attractif en TOPSIS, stable en Monte Carlo, dans un cluster porteur : "
    "la conviction est maximale. La concordance inter-méthodes est notre assurance qualité.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 3 — SYNTHÈSE
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "Partie 3 — Synthèse et Recommandations")

h2(doc, "8.  Top 10 Marchés Attractifs à Horizon 2030")

body(doc,
    "Classement consolidé : Score TOPSIS PCA + Rang Monte Carlo P50 + Cluster + Badge. "
    "Consultable et filtrable en temps réel dans l'application SCAR d'Atlantic Re.")

table(doc,
    ["Rang", "Pays", "Région", "Score\nTOPSIS", "Rang\nMC P50", "Cluster", "Badge"],
    [
        ["1",  "Maroc",         "Afrique du Nord",    "0,78", "1",  "A – Mature",         "ATTRACTIF"],
        ["2",  "Côte d'Ivoire", "Afrique de l'Ouest", "0,71", "2",  "B – Fort potentiel", "ATTRACTIF"],
        ["3",  "Kenya",         "Afrique de l'Est",   "0,68", "3",  "B – Fort potentiel", "ATTRACTIF"],
        ["4",  "Ghana",         "Afrique de l'Ouest", "0,65", "4",  "B – Fort potentiel", "ATTRACTIF"],
        ["5",  "Tunisie",       "Afrique du Nord",    "0,63", "5",  "A – Mature",         "ATTRACTIF"],
        ["6",  "Sénégal",       "Afrique de l'Ouest", "0,60", "6",  "B – Fort potentiel", "ATTRACTIF"],
        ["7",  "Algérie",       "Afrique du Nord",    "0,57", "7",  "A – Mature",         "ATTRACTIF"],
        ["8",  "Égypte",        "Afrique du Nord",    "0,54", "8",  "B – Fort potentiel", "NEUTRE"],
        ["9",  "Botswana",      "Afrique Australe",   "0,52", "9",  "C – Niche",          "NEUTRE"],
        ["10", "Cameroun",      "Afrique Centrale",   "0,49", "11", "B – Fort potentiel", "NEUTRE"],
    ],
    col_widths=[1.0, 2.8, 3.5, 2.0, 2.0, 3.5, 2.2])

body(doc,
    "* Classement indicatif — scores exacts et analyse de sensibilité disponibles dans SCAR.",
    size=11, italic=True, color=GRAY, sa=4)

h2(doc, "9.  Observations Clés")

bullet(doc, "Afrique du Nord (Maroc, Tunisie, Algérie) : cœur de maturité assurantielle — marchés stables et prévisibles.")
bullet(doc, "Afrique de l'Ouest (Côte d'Ivoire, Ghana, Sénégal) : principal moteur de croissance, portée par la démographie et l'urbanisation.")
bullet(doc, "Kenya : hub assurantiel d'Afrique de l'Est, écosystème innovant (mobile insurance).")
bullet(doc, "Égypte : fort potentiel de volume, mais stabilité institutionnelle plus incertaine → badge NEUTRE.")
bullet(doc, "10 marchés ATTRACTIFS sur 33 : concentrer les efforts sur ce cœur maximise le retour sur investissement commercial.")

h2(doc, "10.  Conclusion")

body(doc,
    "Ce travail combine prédiction statistique avancée (Ridge, ARIMA, Conformal Prediction) "
    "et aide à la décision multi-critères (TOPSIS, PCA, Monte Carlo, K-Means) "
    "pour produire un outil d'aide à la décision stratégique robuste, objectif et transparent.")

body(doc,
    "Les résultats sont consultables en temps réel dans l'application SCAR d'Atlantic Re : "
    "visualisations interactives, filtres par région ou badge, simulation de pondérations "
    "personnalisées, et téléchargement des données pays par pays.")

summary_box(doc,
    "Ce rapport est un document vivant : les modèles sont recalibrés annuellement "
    "à chaque mise à jour des données. La méthodologie reste constante, "
    "permettant des comparaisons fiables dans le temps.")


# ══════════════════════════════════════════════════════════════════════════════
# SAUVEGARDE
# ══════════════════════════════════════════════════════════════════════════════
output = r"C:\Users\SMAIKI\AtlanticRe\prediction\Guide_Predictions_Scoring_AtlanticRe_v2.docx"
doc.save(output)
print(f"Rapport généré : {output}")
